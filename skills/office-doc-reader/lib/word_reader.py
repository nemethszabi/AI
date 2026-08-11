"""
Reads an existing .docx document and renders its content as Markdown - the read-side counterpart to
word_helpers.py in the office-doc-builder skill (which only writes/formats). Used to turn inbound client
documents (design docs, RFPs, specs) into text an agent can read and cite.

Mechanical extraction only - no interpretation of what the content means. Paragraphs and tables are
walked in original document order (not paragraphs-then-tables) so a table embedded mid-section stays next
to the text around it. Heading styles ("Heading 1".."Heading 9") become Markdown '#' headings at the
matching level; everything else becomes plain paragraph text.

Known limitation: images, text boxes, and footnotes/endnotes are not extracted - only body paragraphs and
tables. A document that relies on those for meaning will read as incomplete; note that explicitly in the
caller's ingest report rather than silently presenting the extraction as complete.

Requires: python-docx (pip install python-docx).

CLI usage: python word_reader.py <input.docx> <output.md>
"""

import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph


def _iter_block_items(document):
    """Yields Paragraph and Table objects in the order they appear in the document body - python-docx's
    own .paragraphs/.tables accessors return each kind separately, losing interleaving order."""
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def _heading_level(paragraph):
    style_name = (paragraph.style.name or "") if paragraph.style else ""
    if style_name.startswith("Heading "):
        suffix = style_name[len("Heading "):].strip()
        if suffix.isdigit():
            return min(int(suffix), 6)
    if style_name in ("Title",):
        return 1
    return None


def _table_to_markdown(table):
    lines = []
    for i, row in enumerate(table.rows):
        cells = [c.text.strip().replace("|", "\\|").replace("\n", " ") for c in row.cells]
        lines.append("| " + " | ".join(cells) + " |")
        if i == 0:
            lines.append("|" + "---|" * len(cells))
    return "\n".join(lines)


def extract_docx_to_markdown(path):
    """Reads `path` and returns (markdown_text, warnings_list). Empty paragraphs are collapsed (not
    rendered as blank Markdown lines) so the output isn't dominated by whitespace-only lines."""
    document = Document(path)
    warnings = []
    parts = [f"# {Path(path).name}\n"]

    table_count = 0
    paragraph_count = 0

    for block in _iter_block_items(document):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if not text:
                continue
            paragraph_count += 1
            level = _heading_level(block)
            if level:
                parts.append(f"\n{'#' * (level + 1)} {text}\n")
            else:
                parts.append(text + "\n")
        elif isinstance(block, Table):
            table_count += 1
            parts.append("\n" + _table_to_markdown(block) + "\n")

    if paragraph_count == 0 and table_count == 0:
        warnings.append("No extractable text or tables found - document may be image-based or empty")

    return "\n".join(parts), warnings


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("usage: python word_reader.py <input.docx> <output.md>", file=sys.stderr)
        sys.exit(1)

    in_path, out_path = sys.argv[1], sys.argv[2]
    text, warnings = extract_docx_to_markdown(in_path)
    Path(out_path).write_text(text, encoding="utf-8")

    if warnings:
        print("WARNINGS:")
        for w in warnings:
            print(f"  - {w}")
    print(f"Wrote {out_path}")
