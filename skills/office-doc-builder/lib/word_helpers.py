"""
Reusable python-docx formatting helpers - generalized so future Word-generating skills/scripts don't
re-derive the same boilerplate each time.

Requires: python-docx (pip install python-docx). Note the import name is "docx", not "python_docx".
"""

from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_heading_styled(doc, text, level=1):
    """Thin wrapper over doc.add_heading - exists so every heading in a generated document goes through
    one place, in case a consistent style tweak is needed later across many scripts."""
    return doc.add_heading(text, level=level)


def add_toc(doc):
    """Inserts a real Word TOC field (not a static list of headings). IMPORTANT LIMITATION, not a bug:
    Word computes TOC page numbers/entries only when the field is refreshed inside Word itself - python-
    docx cannot pre-render the final content, since that depends on final pagination. The field is
    inserted correctly and will show a placeholder until the user opens the document in Word and does
    Right-click -> Update Field (or press F9, or Word will often prompt automatically on open with
    'This document contains fields that may refer to other files... update?'). Always tell the user this
    when handing over a document that used this helper - don't imply the TOC is already populated.
    """
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    r_element = run._r

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'

    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")

    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click and choose Update Field to generate the table of contents."

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    r_element.append(fld_begin)
    r_element.append(instr_text)
    r_element.append(fld_separate)
    r_element.append(placeholder)
    r_element.append(fld_end)
    return paragraph


def add_table_styled(doc, headers, rows, style="Light Grid Accent 1"):
    """Adds a table with a header row (bold) plus data rows, using one of Word's built-in table styles
    by name (default: a clean grid style available in the default template). headers: list of strings.
    rows: list of lists of cell values (converted to str). Returns the created table."""
    table = doc.add_table(rows=1, cols=len(headers))
    try:
        table.style = style
    except KeyError:
        # Style name not present in this template - fall back to the default rather than raising,
        # since the table content itself still matters more than the exact style name matching.
        pass
    header_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        header_cells[i].text = str(h)
        for p in header_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row_values in rows:
        cells = table.add_row().cells
        for i, v in enumerate(row_values):
            cells[i].text = str(v)
    return table


def add_page_break(doc):
    doc.add_page_break()


def set_default_font(doc, name="Calibri", size_pt=11):
    """Sets the document's Normal style font - affects all text using the default style, applied once
    near the top of a generation script rather than per-run."""
    style = doc.styles["Normal"]
    style.font.name = name
    style.font.size = Pt(size_pt)


def center_paragraph(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
